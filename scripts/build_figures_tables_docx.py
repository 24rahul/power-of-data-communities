#!/usr/bin/env python3
"""
Build two standalone DOCX files for the CCS revision:
  - main_ccs_figures.docx : every figure as an image + a clean text caption
  - main_ccs_tables.docx  : every table as a native Word table + numbered caption

Figure captions are cleaned from LaTeX to plain text via pandoc. Tables are converted
to native Word tables via pandoc (captions get a "Table N" prefix injected first).
"""

import re
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
MS = ROOT / "manuscript_v2"
FIGDIR = MS / "figures"
TEX = (MS / "main_ccs.tex").read_text()


# ---- balanced-brace caption extractor ---------------------------------------
def balanced(s, open_idx):
    depth = 0
    for i in range(open_idx, len(s)):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return s[open_idx + 1:i], i + 1
    return None, len(s)


def all_captions(tex):
    caps = []
    for m in re.finditer(r"\\caption\{", tex):
        content, _ = balanced(tex, m.end() - 1)
        caps.append(content)
    return caps


def latex_to_plain(s):
    """Clean a LaTeX caption/snippet to plain text via pandoc."""
    out = subprocess.run(
        ["pandoc", "-f", "latex", "-t", "plain", "--wrap=none"],
        input=s, capture_output=True, text=True,
    ).stdout.strip()
    return re.sub(r"\s+", " ", out)


CAPS = all_captions(TEX)            # 8 figure caps, then 4 supp-table caps

# ---- FIGURES ----------------------------------------------------------------
# (number label, image path, caption index in CAPS)
FIGURES = [
    ("Figure 1",  "/tmp/fig_concept.png",                       0),
    ("Figure 2",  FIGDIR / "fig1_cumulative_citations.png",     1),
    ("Figure 3",  FIGDIR / "fig5_disciplinary_diffusion.png",   2),
    ("Figure 4",  FIGDIR / "fig3_iceberg_over_time.png",        3),
    ("Figure 5",  FIGDIR / "fig_journal_tiers.png",             4),
    ("Figure 1S", FIGDIR / "fig2_funding_normalized_ripple.png",5),
    ("Figure 2S", FIGDIR / "figure_2_new.png",                  6),
    ("Figure 3S", FIGDIR / "figure_5_second_degree.png",        7),
]


def build_figures(out_path):
    doc = Document()
    doc.core_properties.author = "Milit Patel"
    h = doc.add_heading("Figures", level=1)
    note = doc.add_paragraph(
        "Figures 1–5 are main-text figures; Figures 1S–3S are supplementary.")
    note.runs[0].italic = True

    for n, (label, img, cap_i) in enumerate(FIGURES):
        # image, centered, fit to content width
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img), width=Inches(6.3))
        # caption
        cap = doc.add_paragraph()
        r = cap.add_run(f"{label}. ")
        r.bold = True
        cap.add_run(latex_to_plain(CAPS[cap_i]))
        cap.paragraph_format.space_after = Pt(6)
        if n < len(FIGURES) - 1:
            doc.add_page_break()

    doc.save(out_path)
    print(f"Wrote {out_path} ({len(FIGURES)} figures).")


# ---- TABLES -----------------------------------------------------------------
SUPP_LABELS = ["Table 1S", "Table 2S", "Table 3S", "Table 4S"]
MAIN_LABELS = ["Table 1", "Table 2", "Table 3"]


def caption_as_paragraph(table_tex, label):
    """Move the table's \\caption{...} out to a bold paragraph BEFORE the float, so
    pandoc renders the caption as plain text (no duplicate 'Table N:' auto-number)."""
    m = re.search(r"\\caption\{", table_tex)
    if not m:
        return table_tex
    content, end = balanced(table_tex, m.end() - 1)
    block_wo_caption = table_tex[:m.start()] + table_tex[end:]
    para = f"\\noindent\\textbf{{{label}.}} {content}\\par\\medskip\n"
    return para + block_wo_caption


def build_tables(out_path):
    blocks = []
    for lbl, name in zip(MAIN_LABELS, ["table1", "table2", "table3"]):
        blocks.append(caption_as_paragraph((MS / "tables" / f"{name}.tex").read_text(), lbl))
    supp = re.findall(r"\\begin\{table\}.*?\\end\{table\}", TEX, re.S)
    for lbl, blk in zip(SUPP_LABELS, supp):
        blocks.append(caption_as_paragraph(blk, lbl))

    header = (
        r"\documentclass{article}" "\n"
        r"\usepackage[margin=1in]{geometry}" "\n"
        r"\usepackage{booktabs,array,multirow,amsmath,textcomp}" "\n"
        r"\usepackage{float}" "\n"
        r"\renewcommand{\arraystretch}{1.15}" "\n"
        r"\begin{document}" "\n"
    )
    combined = header + "\n\n".join(blocks) + "\n\\end{document}\n"
    tmp_tex = Path("/tmp/tables_all.tex")
    tmp_tex.write_text(combined)
    subprocess.run(
        ["pandoc", "-f", "latex", "-t", "docx", str(tmp_tex), "-o", str(out_path)],
        check=True,
    )
    # set author property
    doc = Document(out_path)
    doc.core_properties.author = "Milit Patel"
    doc.save(out_path)
    print(f"Wrote {out_path} ({len(blocks)} tables).")


if __name__ == "__main__":
    build_figures(MS / "main_ccs_figures.docx")
    build_tables(MS / "main_ccs_tables.docx")
